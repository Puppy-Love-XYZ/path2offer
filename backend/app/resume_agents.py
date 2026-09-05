

import io
import json
import math
import os
import re
import asyncio
import time
from typing import AsyncGenerator
from dotenv import load_dotenv

load_dotenv(os.path.join(os.path.dirname(__file__), "..", ".env"))

try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage



def _get_llm(model: str = "qwen-plus", temperature: float = 0.1) -> ChatOpenAI:
    api_key = os.getenv("DASHSCOPE_API_KEY") or os.getenv("OPENAI_API_KEY")
    api_base = "https://dashscope.aliyuncs.com/compatible-mode/v1"

    if not api_key:
        raise ValueError("请在 backend/.env 文件中配置 DASHSCOPE_API_KEY（通义千问的API Key）")

    return ChatOpenAI(
        model=model,
        temperature=temperature,
        api_key=api_key,
        base_url=api_base,
        timeout=120,
        max_retries=1,
        max_tokens=4096,
        extra_body={"enable_thinking": False},
    )



_IMAGE_EXTS = {"jpg", "jpeg", "png", "webp", "bmp"}
_IMAGE_MIME = {
    "jpg": "image/jpeg", "jpeg": "image/jpeg",
    "png": "image/png", "webp": "image/webp", "bmp": "image/bmp",
}


def _ocr_image_to_text(file_content: bytes, ext: str) -> str:
   
    import base64, httpx
    api_key = os.getenv("DASHSCOPE_API_KEY") or os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("请配置 DASHSCOPE_API_KEY 以支持图片简历识别")
    mime = _IMAGE_MIME.get(ext, "image/jpeg")
    b64 = base64.b64encode(file_content).decode()
    payload = {
        "model": "qwen-vl-plus",
        "messages": [{
            "role": "user",
            "content": [
                {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                {"type": "text", "text": "请提取图片中简历的所有文字内容，保持原有段落结构输出，不要添加任何解释。"},
            ],
        }],
    }
    resp = httpx.post(
        "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions",
        json=payload,
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        timeout=60,
    )
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"]


def extract_resume_text(file_content: bytes, filename: str) -> str:
   
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "pdf":
        if not HAS_PDF:
            raise RuntimeError("pdfplumber 未安装，请执行: pip install pdfplumber")
        text_parts: list[str] = []
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)
        return "\n".join(text_parts)

    elif ext in ("doc", "docx"):
        if not HAS_DOCX:
            raise RuntimeError("python-docx 未安装，请执行: pip install python-docx")
        doc = docx.Document(io.BytesIO(file_content))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        # 提取文本框内容 (python-docx 不直接支持文本框，需从 XML 中提取)
        try:
            from docx.oxml.ns import qn
            for txbox in doc.element.iter(qn('w:txbxContent')):
                for p in txbox.iter(qn('w:p')):
                    texts = list(p.iter(qn('w:t')))
                    line = "".join(t.text or "" for t in texts)
                    if line.strip():
                        parts.append(line)
        except Exception:
            pass
        return "\n".join(parts)

    elif ext in _IMAGE_EXTS:
        return _ocr_image_to_text(file_content, ext)

    else:
        for enc in ("utf-8", "gbk", "utf-16"):
            try:
                return file_content.decode(enc)
            except Exception:
                continue
        return file_content.decode("utf-8", errors="replace")



_EXTRACT_SYSTEM = """你是一位专业的简历解析专家。
从简历文本中精确提取结构化信息，以 JSON 格式输出。
规则：
- 字段无信息时使用 null 或 []
- 不添加虚假内容，保持原文准确性
- 只输出 JSON，不输出任何解释文字"""

_EXTRACT_HUMAN = """从以下简历原文中提取结构化信息，返回 JSON：

{{
  "basic_info": {{"name":null,"phone":null,"email":null,"location":null,"linkedin":null,"github":null,"target_position":null,"self_intro":null}},
  "education": [{{"school":null,"degree":null,"major":null,"gpa":null,"rank":null,"start_year":null,"end_year":null,"awards":[],"courses":[]}}],
  "work_experience": [{{"company":null,"position":null,"start_date":null,"end_date":null,"description":null,"achievements":[]}}],
  "internships": [{{"company":null,"position":null,"start_date":null,"end_date":null,"description":null,"achievements":[]}}],
  "campus_experience": [{{"organization":null,"role":null,"start_date":null,"end_date":null,"description":null,"scale":null}}],
  "projects": [{{"name":null,"role":null,"tech_stack":[],"start_date":null,"end_date":null,"description":null,"achievements":[],"github":null,"link":null}}],
  "skills": {{"technical":[],"tools":[],"languages":[],"certifications":[],"other":[]}},
  "awards": [],
  "publications": []
}}

简历原文：
{resume_text}"""


_EVAL_SYSTEM = """你是一位拥有15年经验的资深职业顾问，长期服务于互联网、金融、制造等各行业企业的招聘团队及求职者。
精通招聘筛选逻辑：以学历背景为基础、工作/实习经历为核心竞争力、项目成果与技能为亮点。
熟练运用 STAR 原则、量化成就评估方法，能精准识别简历的亮点与短板。

评估原则：
- 结合应聘者实际背景（应届生/有经验求职者），围绕真实招聘市场需求给出评估
- 评分客观不虚高，有据可查
- 建议具体可操作，不说大话套话，聚焦可落地的改进动作
- 只输出 JSON，不输出解释文字"""

_EVAL_HUMAN = """对以下结构化简历数据进行全面深度评估，返回详细 JSON 报告。

简历数据：
{resume_json}

返回以下结构的 JSON（所有分数为整数）：

{{
  "overall_score": 75,
  "overall_rating": "良好",
  "summary": "2-3句总体评价，突出最大亮点和核心不足，结合真实招聘市场需求",
  "dimensions": {{
    "education":         {{"score":80,"rating":"良好","weight":25,"icon":"🎓","label":"学历背景","level":"normal","highlights":[],"issues":[],"suggestions":[],"detail":"详细分析：院校层次/学历/GPA/相关课程/竞赛获奖对求职的影响，结合目标岗位实际招聘门槛"}},
    "internship":        {{"score":70,"rating":"良好","weight":25,"icon":"🏢","label":"工作/实习经历","level":"normal","highlights":[],"issues":[],"suggestions":[],"detail":"详细分析：公司层次/岗位匹配度/成果量化/STAR 表达/对目标职位的加分程度"}},
    "campus_experience": {{"score":65,"rating":"一般","weight":15,"icon":"🏛️","label":"综合经历","level":"normal","highlights":[],"issues":[],"suggestions":[],"detail":"详细分析：学生组织/社团/竞赛/科研/志愿/副业等经历对综合素质和岗位竞争力的体现"}},
    "projects":          {{"score":68,"rating":"一般","weight":15,"icon":"🚀","label":"项目经历","level":"normal","highlights":[],"issues":[],"suggestions":[],"detail":"详细分析：项目相关性/技术深度/个人贡献/成果展示/与目标岗位的匹配程度"}},
    "skills":            {{"score":70,"rating":"良好","weight":10,"icon":"⚡","label":"技能水平","level":"normal","highlights":[],"issues":[],"suggestions":[],"detail":"详细分析：技能与目标岗位的契合度/掌握深度/证书/是否形成差异化优势"}},
    "writing_quality":   {{"score":60,"rating":"一般","weight":5,"icon":"✍️","label":"简历写作","level":"normal","highlights":[],"issues":[],"suggestions":[],"detail":"详细分析：动词力度/量化比例/STAR 结构完整性/篇幅布局/一致性"}},
    "bonus_factors":     {{"score":50,"rating":"待提升","weight":5,"icon":"⭐","label":"加分项","level":"normal","highlights":[],"issues":[],"suggestions":[],"detail":"详细分析：竞赛获奖/论文/开源项目/副业/特殊技能等差异化亮点"}}
  }},
  "strengths": ["最大优势1（具体说明，结合真实岗位竞争力）","优势2","优势3"],
  "weaknesses": ["核心不足1（建设性描述，说明对求职的实际影响）","不足2","不足3"],
  "key_recommendations": [
    {{"priority":"高","category":"内容优化","action":"具体行动（动词开头）","example":"改写前 → 改写后示例","impact":"对简历通过率的预期提升"}},
    {{"priority":"高","category":"经历/项目补充","action":"具体行动","example":"示例","impact":"预期效果"}},
    {{"priority":"中","category":"写作优化","action":"具体行动","example":"改写前 → 改写后示例","impact":"预期效果"}},
    {{"priority":"低","category":"差异化亮点","action":"具体行动","example":"示例","impact":"预期效果"}}
  ],
  "career_path": {{
    "suitable_positions": ["适合岗位方向1","方向2","方向3","方向4"],
    "target_companies": {{"tier1":["冲击（头部/知名大厂）"],"tier2":["稳妥（中大型知名企业）"],"tier3":["保底（成长型中小企业）"]}},
    "short_term": "近期3-6个月具体行动计划（技能提升/项目补充/证书/实习等，结合简历现状给出可落地建议）",
    "medium_term": "1-2年发展规划（结合目标方向和现有基础，给出务实的成长路径）"
  }},
  "interview_tips": ["面试建议1（针对简历弱项或常见追问点）","建议2","建议3"],
  "format_issues": ["简历格式问题（如有）"]
}}

level 字段规则（必须严格执行）：
- "strong"：该维度得分 ≥ 75，显著优势
- "normal"：该维度得分 55-74，正常水平
- "weak"：该维度得分 < 55，明显短板，需重点改进

评分标准（严格执行，不要虚高）：
学历：985本科≥80，211≥70，双非本科50-65，专科≤45；硕士+10，博士+15；GPA前20%+5；相关竞赛+5
工作/实习：无=0；非知名企业1段=35-45；知名企业1段=55-65；头部大厂/顶级机构≥70；多段+10；成果量化清晰+10；与目标岗位强相关+5
综合经历：无=25；普通成员=40；负责人/骨干=55；主要负责人/全国级竞赛奖=70+；顶级荣誉=85+
项目：无=25；课程/学习项目1个=40；个人独立项目有成果=55；团队核心贡献有落地=65-75；开源/线上产品/论文=80+
技能：简单罗列=35；基础掌握=50；熟练+证书=65；精通+权威证书=80+；竞赛/实战验证=85+
写作：泛泛描述=35；有少量量化=50；50%以上量化=65；全量化+STAR完整=80+
加分：无差异化=25；一般奖项=45；省级/全国=65；顶会论文/知名开源=85+
综合加权后根据整体市场竞争力微调±5分。"""


_JD_MATCH_SYSTEM = """你是一位拥有10年经验的资深招聘顾问，擅长将候选人简历与具体岗位要求进行精准匹配分析。
评估视角：
- 结合候选人实际背景（应届/有经验），合理判断能力与岗位的契合度
- High Match：简历中有明确的经历/项目/技能直接支撑 JD 要求
- Potential Match：有相关背景但描述不够直接，有潜力但需要包装或补充
- Gap：JD 要求的核心技能/经历在简历中完全缺失
评估原则：精准客观、具体可操作、贴近真实招聘场景。只输出 JSON，不输出解释文字。"""

_JD_MATCH_HUMAN = """基于以下岗位描述（JD），对候选人简历进行逐条精确匹配分析，返回详细 JSON 报告。
岗位描述（JD）：
{jd_text}
结构化简历数据：
{resume_json}
返回以下 JSON 结构（所有字段必填）：
{{
  "overall_match": "High|Medium|Low",
  "match_percentage": 75,
  "summary": "2-3句精准判断：候选人与该岗位的核心契合点和主要差距，结合实际招聘标准",
  "requirements_analysis": [
    {{
      "requirement": "从JD提取的具体要求（5-8条，含技术栈/学历/经历/软技能等）",
      "match_level": "High Match|Potential Match|Gap",
      "evidence": "简历中对应的具体证据（High/Potential时引用简历原文）或缺失说明（Gap时）",
      "rewrite_suggestion": "具体改写建议（Gap和Potential必须给出'改写前 → 改写后'可直接使用的示例，High时可提供强化建议）"
    }}
  ],
  "strengths": ["相对该JD的核心优势1（引用简历具体内容，说明对岗位竞争力的加分）","优势2","优势3"],
  "gaps": ["核心缺失1（说明对通过简历筛选的影响程度）","缺失2","缺失3"],
  "rewrite_priority": [
    {{
      "priority": "高|中|低",
      "section": "简历模块（如：工作/实习经历/项目经历/技能/自我介绍）",
      "action": "具体改写动作（动词开头，针对性强）",
      "example": "改写前：原文内容 → 改写后：优化后内容（与JD关键词对齐，量化成果）"
    }}
  ],
  "interview_prediction": {{
    "pass_rate": "高|中|低",
    "likely_questions": [
      "面试官最可能追问的问题1（基于JD重点和简历弱项）",
      "问题2",
      "问题3"
    ],
    "red_flags": ["面试官可能质疑的点1（如：经历时间短/项目参与度不明确/技能描述模糊）","质疑点2"]
  }},
  "final_verdict": "综合建议（强烈推荐投递/建议投递/谨慎投递/不建议投递）及核心理由（1-2句）"
}}

分析要求：
1. 从JD中提取5-8个核心要求，包含技术栈、学历要求、工作/实习/项目经验、软技能等
2. 对每条要求按三级标准客观判定，匹配程度要贴近真实招聘标准
3. 改写建议具体到句子，给出可直接使用的"改写前 → 改写后"示例，注意与JD关键词对齐
4. interview_prediction 聚焦实际面试高频考点：项目深挖、经历追问、技术基础考察
5. Gap 的判定要合理，结合候选人背景（应届/有经验）综合评判"""

def _parse_json(content: str) -> dict:
    """从 LLM 输出中解析 JSON，支持 markdown 代码块包裹。"""
    content = content.strip()
    for pattern in [r"```json\s*([\s\S]+?)\s*```", r"```\s*([\s\S]+?)\s*```", r"(\{[\s\S]+\})"]:
        m = re.search(pattern, content, re.DOTALL)
        if m:
            try:
                return json.loads(m.group(1))
            except Exception:
                continue
    try:
        return json.loads(content)
    except Exception as e:
        raise ValueError(f"无法解析 JSON：{content[:300]}…\n错误：{e}")


def _sse(data: dict) -> str:
    return f"data: {json.dumps(data, ensure_ascii=False)}\n\n"


def _fmt_err(e: Exception) -> str:
  
    msg = str(e).replace("\n", " ").strip()
    if "Connection error" in msg or "ConnectError" in msg:
        return "网络连接失败，请检查网络或 DASHSCOPE_API_KEY 是否正确"
    if "401" in msg or "Unauthorized" in msg or "Invalid" in msg:
        return "API Key 无效或已过期，请检查 DASHSCOPE_API_KEY"
    if "404" in msg:
        return "API 地址错误，请确认 base_url 配置"
    if "429" in msg or "rate limit" in msg.lower():
        return "请求频率超限（Rate Limit），稍后重试"
    if "timeout" in msg.lower():
        return "请求超时，请检查网络或稍后重试"
    if "inappropriate content" in msg.lower() or "DataInspection" in msg or "content_filter" in msg.lower():
        return "内容安全审查未通过（阿里云内容审核），已自动重试仍失败，请更换简历内容后再试"
    return msg[:200]


def _is_content_filter_error(e: Exception) -> bool:
    msg = str(e)
    return (
        "inappropriate content" in msg.lower()
        or "DataInspection" in msg
        or "content_filter" in msg.lower()
    )


async def _stream_llm(
    llm: ChatOpenAI,
    messages: list,
    step: int,
    stage: str,
    message: str,
    progress_start: float,
    progress_end: float,
) -> AsyncGenerator[str, None]:
    content_parts: list[str] = []
    start_t = time.time()
    last_hb = 0.0
    
    tau = 12.0
    span = progress_end - progress_start - 3  

    async for chunk in llm.astream(messages):
        if chunk.content:
            content_parts.append(chunk.content)

        now = time.time()
        if now - last_hb >= 0.4:
            elapsed = now - start_t
            ratio = 1.0 - math.exp(-elapsed / tau)
            prog = int(progress_start + span * ratio)
            yield _sse({"step": step, "stage": stage, "message": message, "progress": prog})
            last_hb = now

    yield _sse({"__content__": "".join(content_parts)})



async def extract_resume_structured(file_content: bytes, filename: str) -> dict:

    from langchain_core.messages import HumanMessage, SystemMessage

    raw_text = extract_resume_text(file_content, filename)
    if not raw_text.strip():
        raise ValueError("未能从文件中提取到文本，请检查文件格式")

    llm = _get_llm("qwen-plus", temperature=0.05)
    resp = await llm.ainvoke([
        SystemMessage(content=_EXTRACT_SYSTEM),
        HumanMessage(content=_EXTRACT_HUMAN.format(resume_text=raw_text[:8000])),
    ])
    return _parse_json(resp.content)



async def run_evaluation_pipeline(
    file_content: bytes,
    filename: str,
) -> AsyncGenerator[str, None]:
    pipeline_start = time.time()
 # ── Step 1: 读取文件 ────────────────────────────────────────────────────
    yield _sse({"step": 1, "total": 5, "stage": "extract", "message": "📄 正在读取简历文件...", "progress": 5})

    try:
        raw_text = extract_resume_text(file_content, filename)
    except Exception as e:
        yield _sse({"step": 1, "total": 5, "stage": "error",
                    "message": f"文件读取失败：{e}", "progress": 5, "error": True})
        return

    if not raw_text.strip():
        yield _sse({"step": 1, "total": 5, "stage": "error",
                    "message": "⚠️ 未能提取到文本内容，请检查文件格式", "progress": 5, "error": True})
        return

    char_count = len(raw_text)
    yield _sse({"step": 1, "total": 5, "stage": "extract",
                "message": f"✅ 文件读取完成（{char_count} 字符）", "progress": 15})

    
    yield _sse({"step": 2, "total": 5, "stage": "parse",
                "message": "🔍 Agent-1 正在解析简历结构...", "progress": 18})

    extract_content = ""
    _extract_msgs = [
        SystemMessage(content=_EXTRACT_SYSTEM),
        HumanMessage(content=_EXTRACT_HUMAN.format(resume_text=raw_text[:8000])),
    ]
    for attempt in range(2):
        extract_content = ""
        try:
            llm_extract = _get_llm("qwen-plus", 0.05)
            async for event in _stream_llm(
                llm_extract, _extract_msgs,
                step=2, stage="parse",
                message="🔍 Agent-1 正在解析简历结构...",
                progress_start=18, progress_end=45,
            ):
                data = json.loads(event[len("data: "):].strip())
                if "__content__" in data:
                    extract_content = data["__content__"]
                else:
                    yield event
            break  # 成功
        except Exception as e:
            if attempt == 0 and _is_content_filter_error(e):
                yield _sse({"step": 2, "total": 5, "stage": "parse",
                            "message": "⚠️ 内容安全拦截，正在自动重试...", "progress": 18})
                await asyncio.sleep(1)
                continue
            yield _sse({"step": 2, "total": 5, "stage": "error",
                        "message": f"结构解析失败：{_fmt_err(e)}", "progress": 18, "error": True})
            return
    else:
        yield _sse({"step": 2, "total": 5, "stage": "error",
                    "message": "结构解析失败：内容安全审查拦截，请更换简历文件后重试", "progress": 18, "error": True})
        return

    try:
        resume_data = _parse_json(extract_content)
        yield _sse({"step": 2, "total": 5, "stage": "parse",
                    "message": "✅ 简历结构解析完成", "progress": 45,
                    "parsed_data": resume_data})
    except Exception as e:
        yield _sse({"step": 2, "total": 5, "stage": "error",
                    "message": f"结构解析失败：{_fmt_err(e)}", "progress": 18, "error": True})
        return

    
    yield _sse({"step": 3, "total": 5, "stage": "evaluate",
                "message": "🧠 Agent-2 正在进行多维度深度评估...", "progress": 48})

    _eval_msgs = [
        SystemMessage(content=_EVAL_SYSTEM),
        HumanMessage(content=_EVAL_HUMAN.format(
            resume_json=json.dumps(resume_data, ensure_ascii=False, indent=2)
        )),
    ]
    eval_content = ""
    for attempt in range(2):
        eval_content = ""
        try:
            llm_eval = _get_llm("qwen-plus", 0.2 if attempt == 0 else 0.1)
            async for event in _stream_llm(
                llm_eval, _eval_msgs,
                step=3, stage="evaluate",
                message="🧠 Agent-2 正在进行多维度深度评估...",
                progress_start=48, progress_end=88,
            ):
                data = json.loads(event[len("data: "):].strip())
                if "__content__" in data:
                    eval_content = data["__content__"]
                else:
                    yield event
            break  # 成功
        except Exception as e:
            if attempt == 0 and _is_content_filter_error(e):
                yield _sse({"step": 3, "total": 5, "stage": "evaluate",
                            "message": "⚠️ 内容安全拦截，正在自动重试...", "progress": 48})
                await asyncio.sleep(1)
                continue
            yield _sse({"step": 3, "total": 5, "stage": "error",
                        "message": f"评估失败：{_fmt_err(e)}", "progress": 48, "error": True})
            return
    else:
        yield _sse({"step": 3, "total": 5, "stage": "error",
                    "message": "评估失败：内容安全审查拦截，请更换简历文件后重试", "progress": 48, "error": True})
        return

    try:
        evaluation = _parse_json(eval_content)
        yield _sse({"step": 3, "total": 5, "stage": "evaluate",
                    "message": "✅ 多维度评估完成", "progress": 88})
    except Exception as e:
        yield _sse({"step": 3, "total": 5, "stage": "error",
                    "message": f"评估失败：{_fmt_err(e)}", "progress": 48, "error": True})
        return

    # ── Step 4-5: 汇总报告 ─────────────────────────────────────────────────
    elapsed = round(time.time() - pipeline_start, 1)
    yield _sse({"step": 4, "total": 5, "stage": "report",
                "message": f"📊 正在生成完整分析报告（共耗时 {elapsed}s）...", "progress": 92})

    final_report = {
        "resume_data": resume_data,
        "evaluation": evaluation,
        "filename": filename,
        "char_count": char_count,
        "elapsed_seconds": elapsed,
    }

    yield _sse({
        "step": 5,
        "total": 5,
        "stage": "done",
        "message": f"🎉 评估完成！（总耗时 {elapsed}s）",
        "progress": 100,
        "result": final_report,
    })



async def run_jd_evaluation_pipeline(
    file_content: bytes,
    filename: str,
    jd_text: str,
) -> AsyncGenerator[str, None]:
   
    pipeline_start = time.time()
    TOTAL = 6

    
    yield _sse({"step": 1, "total": TOTAL, "stage": "extract", "message": "📄 正在读取简历文件...", "progress": 5})

    try:
        raw_text = extract_resume_text(file_content, filename)
    except Exception as e:
        yield _sse({"step": 1, "total": TOTAL, "stage": "error",
                    "message": f"文件读取失败：{e}", "progress": 5, "error": True})
        return

    if not raw_text.strip():
        yield _sse({"step": 1, "total": TOTAL, "stage": "error",
                    "message": "⚠️ 未能提取到文本内容，请检查文件格式", "progress": 5, "error": True})
        return

    char_count = len(raw_text)
    yield _sse({"step": 1, "total": TOTAL, "stage": "extract",
                "message": f"✅ 文件读取完成（{char_count} 字符）", "progress": 12})

    
    yield _sse({"step": 2, "total": TOTAL, "stage": "parse",
                "message": "🔍 Agent-1 正在解析简历结构...", "progress": 15})

    _extract_msgs2 = [
        SystemMessage(content=_EXTRACT_SYSTEM),
        HumanMessage(content=_EXTRACT_HUMAN.format(resume_text=raw_text[:8000])),
    ]
    extract_content = ""
    for attempt in range(2):
        extract_content = ""
        try:
            llm_extract = _get_llm("qwen-plus", 0.05)
            async for event in _stream_llm(
                llm_extract, _extract_msgs2,
                step=2, stage="parse",
                message="🔍 Agent-1 正在解析简历结构...",
                progress_start=15, progress_end=40,
            ):
                data = json.loads(event[len("data: "):].strip())
                if "__content__" in data:
                    extract_content = data["__content__"]
                else:
                    yield event
            break
        except Exception as e:
            if attempt == 0 and _is_content_filter_error(e):
                yield _sse({"step": 2, "total": TOTAL, "stage": "parse",
                            "message": "⚠️ 内容安全拦截，正在自动重试...", "progress": 15})
                await asyncio.sleep(1)
                continue
            yield _sse({"step": 2, "total": TOTAL, "stage": "error",
                        "message": f"结构解析失败：{_fmt_err(e)}", "progress": 15, "error": True})
            return
    else:
        yield _sse({"step": 2, "total": TOTAL, "stage": "error",
                    "message": "结构解析失败：内容安全审查拦截，请更换简历文件后重试", "progress": 15, "error": True})
        return

    try:
        resume_data = _parse_json(extract_content)
        yield _sse({"step": 2, "total": TOTAL, "stage": "parse",
                    "message": "✅ 简历结构解析完成", "progress": 40,
                    "parsed_data": resume_data})
    except Exception as e:
        yield _sse({"step": 2, "total": TOTAL, "stage": "error",
                    "message": f"结构解析失败：{_fmt_err(e)}", "progress": 15, "error": True})
        return


    yield _sse({"step": 3, "total": TOTAL, "stage": "evaluate",
                "message": "🧠 Agent-2 正在进行综合多维度评估...", "progress": 43})

    _eval_msgs2 = [
        SystemMessage(content=_EVAL_SYSTEM),
        HumanMessage(content=_EVAL_HUMAN.format(
            resume_json=json.dumps(resume_data, ensure_ascii=False, indent=2)
        )),
    ]
    eval_content = ""
    for attempt in range(2):
        eval_content = ""
        try:
            llm_eval = _get_llm("qwen-plus", 0.2 if attempt == 0 else 0.1)
            async for event in _stream_llm(
                llm_eval, _eval_msgs2,
                step=3, stage="evaluate",
                message="🧠 Agent-2 正在进行综合多维度评估...",
                progress_start=43, progress_end=68,
            ):
                data = json.loads(event[len("data: "):].strip())
                if "__content__" in data:
                    eval_content = data["__content__"]
                else:
                    yield event
            break
        except Exception as e:
            if attempt == 0 and _is_content_filter_error(e):
                yield _sse({"step": 3, "total": TOTAL, "stage": "evaluate",
                            "message": "⚠️ 内容安全拦截，正在自动重试...", "progress": 43})
                await asyncio.sleep(1)
                continue
            yield _sse({"step": 3, "total": TOTAL, "stage": "error",
                        "message": f"综合评估失败：{_fmt_err(e)}", "progress": 43, "error": True})
            return
    else:
        yield _sse({"step": 3, "total": TOTAL, "stage": "error",
                    "message": "综合评估失败：内容安全审查拦截，请更换简历文件后重试", "progress": 43, "error": True})
        return

    try:
        evaluation = _parse_json(eval_content)
        yield _sse({"step": 3, "total": TOTAL, "stage": "evaluate",
                    "message": "✅ 综合评估完成", "progress": 68})
    except Exception as e:
        yield _sse({"step": 3, "total": TOTAL, "stage": "error",
                    "message": f"综合评估失败：{_fmt_err(e)}", "progress": 43, "error": True})
        return

    yield _sse({"step": 4, "total": TOTAL, "stage": "jd_match",
                "message": "🎯 Agent-3 正在进行 JD 精确匹配分析...", "progress": 71})

    _jd_msgs = [
        SystemMessage(content=_JD_MATCH_SYSTEM),
        HumanMessage(content=_JD_MATCH_HUMAN.format(
            jd_text=jd_text[:4000],
            resume_json=json.dumps(resume_data, ensure_ascii=False, indent=2),
        )),
    ]
    jd_content = ""
    for attempt in range(2):
        jd_content = ""
        try:
            llm_jd = _get_llm("qwen-plus", 0.1 if attempt == 0 else 0.05)
            async for event in _stream_llm(
                llm_jd, _jd_msgs,
                step=4, stage="jd_match",
                message="🎯 Agent-3 正在进行 JD 精确匹配分析...",
                progress_start=71, progress_end=88,
            ):
                data = json.loads(event[len("data: "):].strip())
                if "__content__" in data:
                    jd_content = data["__content__"]
                else:
                    yield event
            break
        except Exception as e:
            if attempt == 0 and _is_content_filter_error(e):
                yield _sse({"step": 4, "total": TOTAL, "stage": "jd_match",
                            "message": "⚠️ 内容安全拦截，正在自动重试...", "progress": 71})
                await asyncio.sleep(1)
                continue
            yield _sse({"step": 4, "total": TOTAL, "stage": "error",
                        "message": f"JD匹配分析失败：{_fmt_err(e)}", "progress": 71, "error": True})
            return
    else:
        yield _sse({"step": 4, "total": TOTAL, "stage": "error",
                    "message": "JD匹配分析失败：内容安全审查拦截，请更换简历文件后重试", "progress": 71, "error": True})
        return

    try:
        jd_match = _parse_json(jd_content)
        yield _sse({"step": 4, "total": TOTAL, "stage": "jd_match",
                    "message": "✅ JD 精确匹配分析完成", "progress": 88})
    except Exception as e:
        yield _sse({"step": 4, "total": TOTAL, "stage": "error",
                    "message": f"JD匹配分析失败：{_fmt_err(e)}", "progress": 71, "error": True})
        return

  
    elapsed = round(time.time() - pipeline_start, 1)
    yield _sse({"step": 5, "total": TOTAL, "stage": "report",
                "message": f"📊 正在生成精确匹配报告（共耗时 {elapsed}s）...", "progress": 92})

    final_report = {
        "resume_data": resume_data,
        "evaluation": evaluation,
        "jd_match": jd_match,
        "filename": filename,
        "char_count": char_count,
        "elapsed_seconds": elapsed,
    }

    yield _sse({
        "step": 6,
        "total": TOTAL,
        "stage": "done",
        "message": f"🎉 精确分析完成！（总耗时 {elapsed}s）",
        "progress": 100,
        "result": final_report,
    })

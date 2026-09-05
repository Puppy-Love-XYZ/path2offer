

import os
import re
import threading
import logging
from pathlib import Path
from typing import List, Dict, Optional

import numpy as np

logger = logging.getLogger(__name__)

os.environ.setdefault("HF_ENDPOINT", "https://hf-mirror.com")


_HERE = Path(__file__).parent
_CHROMA_PATH = _HERE / "chroma_db"


MODEL_NAME = "BAAI/bge-small-zh-v1.5"
COLLECTION_NAME = "jobs_v1"


W2V_MODEL_PATH = os.getenv("W2V_MODEL_PATH", "")


EDU_RANK = {"不限": 0, "大专": 1, "本科": 2, "硕士": 3, "博士": 4}


_model = None
_chroma_client = None
_collection = None
_w2v_model = None
_w2v_available = None   

_index_status: dict = {
    "status": "not_started",  
    "total": 0,
    "indexed": 0,
    "error": None,
}
_lock = threading.Lock()



def get_model():
    global _model
    if _model is None:
        from sentence_transformers import SentenceTransformer
        try:
            logger.info("从本地缓存加载 SBERT 模型: %s", MODEL_NAME)
            _model = SentenceTransformer(MODEL_NAME, local_files_only=True)
        except Exception:
            logger.info("本地缓存不存在，尝试从网络下载 SBERT 模型: %s", MODEL_NAME)
            _model = SentenceTransformer(MODEL_NAME)
        logger.info("SBERT 模型加载完成")
    return _model



def get_w2v_model():
    global _w2v_model, _w2v_available
    if _w2v_available is not None:
        return _w2v_model
    if not W2V_MODEL_PATH or not Path(W2V_MODEL_PATH).exists():
        _w2v_available = False
        logger.warning("W2V_MODEL_PATH 未配置或文件不存在，Word2Vec 信号不可用，将降级为纯 SBERT")
        return None
    try:
        from gensim.models import KeyedVectors
        logger.info("加载 Word2Vec 模型: %s", W2V_MODEL_PATH)
        _w2v_model = KeyedVectors.load_word2vec_format(W2V_MODEL_PATH, binary=True)
        _w2v_available = True
        logger.info("Word2Vec 模型加载完成，词汇量: %d", len(_w2v_model))
    except Exception as e:
        _w2v_available = False
        logger.warning("Word2Vec 模型加载失败（%s），降级为纯 SBERT", e)
    return _w2v_model


def _create_chroma_collection():

    global _chroma_client, _collection
    import chromadb
    from chromadb.config import Settings
    _CHROMA_PATH.mkdir(exist_ok=True)
    _chroma_client = chromadb.PersistentClient(
        path=str(_CHROMA_PATH),
        settings=Settings(anonymized_telemetry=False),
    )
    _collection = _chroma_client.get_or_create_collection(
        name=COLLECTION_NAME,
        metadata={"hnsw:space": "cosine"},
    )
    return _collection


def get_collection():

    global _collection
    if _collection is None:
        return _create_chroma_collection()

    try:
        _collection.count()
    except Exception as e:
        if "closed" in str(e).lower():
            logger.warning("ChromaDB 客户端已关闭，正在重建…")
            return _create_chroma_collection()
        raise
    return _collection




def build_job_text(job) -> str:

    parts = []
    if job.job_name:
        parts.append(f"职位：{job.job_name}")
    if job.industry_name:
        parts.append(f"行业：{job.industry_name}")
    if job.company_name:
        parts.append(f"公司：{job.company_name}")
    if job.work_city:
        parts.append(f"城市：{job.work_city}")
    if job.your_education:
        parts.append(f"学历要求：{job.your_education}")
    if job.working_exp:
        parts.append(f"工作经验：{job.working_exp}")
    if job.company_size:
        parts.append(f"公司规模：{job.company_size}")
    if job.job_summary:
        parts.append(f"岗位描述：{job.job_summary[:800]}")
    return "\n".join(parts)


def build_resume_text(parsed: dict) -> str:

    parts = []
    basic = parsed.get("basic_info") or {}


    highest_edu = ""
    for edu in (parsed.get("education") or []):
        degree = edu.get("degree", "")
        major = edu.get("major", "")
        school = edu.get("school", "")
        if degree:
            highest_edu = degree
        line = " ".join(t for t in [school, degree, major] if t)
        if line:
            parts.append(f"教育背景：{line}")
    if highest_edu:
        parts.append(f"最高学历：{highest_edu}")

    if basic.get("major"):
        parts.append(f"专业：{basic['major']}")

    # 工作经历
    for exp in (parsed.get("work_experience") or []):
        company = exp.get("company", "")
        pos = exp.get("position", "")
        desc = exp.get("description", "")[:200]
        parts.append(f"工作经历：{company} {pos} {desc}".strip())

    # 实习经历
    for intern in (parsed.get("internship") or []):
        company = intern.get("company", "")
        pos = intern.get("position", "")
        desc = intern.get("description", "")[:200]
        if company or pos:
            parts.append(f"实习经历：{company} {pos} {desc}".strip())


    for proj in (parsed.get("projects") or parsed.get("project") or []):
        name = proj.get("name", "")
        desc = proj.get("description", "")[:200]
        if name:
            proj_line = f"项目经历：{name} {desc}".strip()
            for _ in range(3):
                parts.append(proj_line)

    skills = parsed.get("skills")
    if isinstance(skills, list) and skills:
        skill_str = "、".join(str(s) for s in skills[:30])
        for _ in range(3):
            parts.append(f"技能：{skill_str}")
    elif isinstance(skills, str) and skills:
        for _ in range(3):
            parts.append(f"技能：{skills[:400]}")

    self_eval = parsed.get("self_evaluation") or parsed.get("summary") or ""
    if self_eval:
        parts.append(f"自我评价：{str(self_eval)[:300]}")

    return "\n".join(parts) if parts else str(parsed)[:600]


_IMAGE_EXTS = {"jpg", "jpeg", "png", "webp", "bmp"}
_IMAGE_MIME = {
    "jpg": "image/jpeg", "jpeg": "image/jpeg",
    "png": "image/png", "webp": "image/webp", "bmp": "image/bmp",
}


def _ocr_image_to_text(content: bytes, ext: str) -> str:

    import base64, httpx
    api_key = os.getenv("DASHSCOPE_API_KEY") or os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("请配置 DASHSCOPE_API_KEY 以支持图片简历识别")
    mime = _IMAGE_MIME.get(ext, "image/jpeg")
    b64 = base64.b64encode(content).decode()
    payload = {
        "model": "qwen-vl-plus",
        "messages": [{
            "role": "user",
            "content": [
                {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                {"type": "text", "text": "请提取图片中简历的所有文字内容，保持原有段落结构输出，不要添加任何解释。"},
            ],
        }],
    }
    resp = httpx.post(
        "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions",
        json=payload,
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        timeout=60,
    )
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"]


def extract_text_from_file(content: bytes, filename: str) -> str:
    """从上传的 PDF / DOCX / TXT / 图片 提取纯文本"""
    fname = (filename or "").lower()
    ext = fname.rsplit(".", 1)[-1] if "." in fname else ""

    if fname.endswith(".pdf"):
        import pdfplumber, io
        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n".join(text_parts)
    elif fname.endswith(".docx"):
        import docx, io
        doc = docx.Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # 表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        # 提取文本框内容 (python-docx 不直接支持文本框，需从 XML 中提取)
        try:
            from docx.oxml.ns import qn
            for txbox in doc.element.iter(qn('w:txbxContent')):
                for p in txbox.iter(qn('w:p')):
                    texts = list(p.iter(qn('w:t')))
                    line = "".join(t.text or "" for t in texts)
                    if line.strip():
                        parts.append(line)
        except Exception:
            pass
        return "\n".join(parts)
    elif ext in _IMAGE_EXTS:
        return _ocr_image_to_text(content, ext)
    else:
        for enc in ("utf-8", "gbk", "utf-16"):
            try:
                return content.decode(enc)
            except Exception:
                continue
        return content.decode("utf-8", errors="ignore")




def _text_to_w2v_vec(text: str, model) -> Optional[np.ndarray]:

    import jieba.analyse
    keywords = jieba.analyse.extract_tags(text, topK=30)
    vecs = [model[w] for w in keywords if w in model]
    if not vecs:
        return None
    v = np.mean(vecs, axis=0)
    norm = np.linalg.norm(v)
    if norm < 1e-8:
        return None
    return v / norm


def _extract_keywords_set(text: str, top_k: int = 30) -> set:

    import jieba.analyse
    return set(jieba.analyse.extract_tags(text, topK=top_k))


def _w2v_expand_keywords(keywords: set, model, top_n: int = 5, threshold: float = 0.65) -> set:

    expanded = set(keywords)
    for kw in keywords:
        if kw not in model:
            continue
        try:
            for word, sim in model.most_similar(kw, topn=top_n):
                if sim >= threshold:
                    expanded.add(word)
        except Exception:
            pass
    return expanded


def _skill_coverage_with_w2v(resume_kws: set, job_kws: set, w2v=None) -> float:

    if not job_kws:
        return 0.0
    if w2v is not None:
        r_exp = _w2v_expand_keywords(resume_kws, w2v)
        j_exp = _w2v_expand_keywords(job_kws, w2v)
    else:
        r_exp, j_exp = resume_kws, job_kws
    return len(r_exp & j_exp) / max(len(j_exp), 1)


def _edu_score(resume_text: str, job_edu_req: str) -> tuple:

    resume_edu = "不限"
    for edu_name in ["博士", "硕士", "本科", "大专"]:
        if edu_name in resume_text:
            resume_edu = edu_name
            break

    req = job_edu_req.strip() if job_edu_req else "不限"

    req_norm = req.replace("学历不限", "不限").replace("大学本科", "本科")
    for k in EDU_RANK:
        if k in req_norm:
            req_norm = k
            break
    else:
        req_norm = "不限"

    r_rank = EDU_RANK.get(resume_edu, 2)   
    j_rank = EDU_RANK.get(req_norm, 0)     

    diff = r_rank - j_rank
    if diff >= 0:
        if diff == 0:
            score, desc = 100, "学历完全匹配"
        elif diff == 1:
            score, desc = 90, "学历略高于要求"
        else:
            score, desc = 85, "学历明显高于要求"
        return score, desc, True
    else:

        if diff == -1:
            score, desc = 60, "学历略低于要求"
        elif diff == -2:
            score, desc = 20, "学历明显低于要求"
        else:
            score, desc = 0, "学历差距较大"
        return score, desc, False


def _exp_score(resume_text: str, job_exp_req: str) -> tuple:


    resume_years = 0
    m = re.search(r"(\d+)\s*年.*?经[验历]", resume_text)
    if m:
        resume_years = int(m.group(1))
    elif "实习" in resume_text:
        resume_years = 0

    req = (job_exp_req or "").strip()
    if not req or req in ("不限", "经验不限", "在校生/应届生"):
        return 100, "经验不限"

    req_min, req_max = 0, 99
    m2 = re.search(r"(\d+)[^\d]*[-~至](\d+)\s*年", req)
    m3 = re.search(r"(\d+)\s*年以上", req)
    m4 = re.search(r"(\d+)\s*年以下", req)
    m5 = re.search(r"(\d+)\s*年", req)
    if m2:
        req_min, req_max = int(m2.group(1)), int(m2.group(2))
    elif m3:
        req_min = int(m3.group(1))
    elif m4:
        req_max = int(m4.group(1))
    elif m5:
        req_min = int(m5.group(1))
        req_max = req_min + 2

    if req_min <= resume_years <= req_max:
        return 100, f"经验符合（要求{req}，实际{resume_years}年）"
    elif resume_years > req_max:
        return 85, f"经验高于要求（要求{req}，实际{resume_years}年）"
    elif resume_years < req_min:
        gap = req_min - resume_years
        if gap <= 1:
            return 60, f"经验略不足（要求{req}，实际{resume_years}年）"
        else:
            return 30, f"经验明显不足（要求{req}，实际{resume_years}年）"
    return 80, f"经验基本匹配（{req}）"




def compute_match_report(resume_text: str, job) -> dict:
 
    model = get_model()
    w2v = get_w2v_model()
    w2v_available = (_w2v_available is True)

    job_text = build_job_text(job)


    vecs = model.encode([resume_text, job_text], normalize_embeddings=True)
    sbert_sim = float(np.dot(vecs[0], vecs[1]))
    sbert_score = round(max(0.0, sbert_sim) * 100, 1)

    r_kw = _extract_keywords_set(resume_text, 30)
    j_kw = _extract_keywords_set(job_text, 20)
    raw_coverage = _skill_coverage_with_w2v(r_kw, j_kw, w2v if w2v_available else None)

    skill_score = round(min(raw_coverage * 1.5 * 100, 100.0), 1)
    matched_skills = sorted(r_kw & j_kw)   #


    exp_req = getattr(job, "working_exp", "") or ""
    exp_score_val, exp_desc = _exp_score(resume_text, exp_req)


    edu_req = getattr(job, "your_education", "") or ""
    edu_score, edu_desc, edu_meets = _edu_score(resume_text, edu_req)


    total_raw = 0.40 * sbert_score + 0.40 * skill_score + 0.20 * exp_score_val
    if not edu_meets:

        penalty = 0.80 if edu_score >= 60 else 0.65
        total_raw *= penalty
    total = round(max(0.0, min(100.0, total_raw)), 1)


    highlights: list[str] = []
    suggestions: list[str] = []

    if sbert_score >= 70:
        highlights.append(f"语义匹配度高（{sbert_score}分），简历方向与岗位高度吻合")
    elif sbert_score < 50:
        suggestions.append("简历与岗位语义差距较大，建议突出与岗位相关的经历和描述")

    if skill_score >= 60:
        w2v_note = "（已启用同义词扩展）" if w2v_available else ""
        highlights.append(f"技能关键词匹配良好（{skill_score}分）{w2v_note}")
    elif skill_score < 30:
        missing = sorted(j_kw - r_kw)[:5]
        tip = f"：{', '.join(missing)}" if missing else ""
        suggestions.append(f"简历缺少岗位核心技能关键词{tip}，建议针对 JD 优化表述")

    if matched_skills:
        highlights.append(f"命中技能词：{', '.join(matched_skills[:5])}")

    if not edu_meets:
        penalty_str = "0.80" if edu_score >= 60 else "0.65"
        suggestions.append(f"学历未达门槛（{edu_desc}），总分已施加 × {penalty_str} 折扣")
    elif edu_score >= 90:
        highlights.append(f"学历满足要求（{edu_desc}）")

    if exp_score_val >= 90:
        highlights.append(f"经验符合要求（{exp_desc}）")
    elif exp_score_val < 50:
        suggestions.append(f"经验不足：{exp_desc}，可突出实习/项目经历加以弥补")

    return {
        "total_score": total,
        "sbert_score": sbert_score,
        "w2v_score": None,           # 旧平均向量法已废弃
        "w2v_available": w2v_available,
        "edu_meets_threshold": edu_meets,
        "dimensions": {
            "semantic": {
                "score": sbert_score,
                "weight": 0.40,
                "label": "语义理解",
            },
            "skills": {
                "score": skill_score,
                "weight": 0.40,
                "label": "技能匹配",
                "matched_skills": matched_skills[:10],
                "matched_count": len(matched_skills),
                "total_count": len(j_kw),
                "w2v_used": w2v_available,
            },
            "experience": {
                "score": exp_score_val,
                "weight": 0.20,
                "match": exp_desc,
                "label": "经验修正",
            },
            "education": {
                "score": edu_score,
                "meets_threshold": edu_meets,
                "match": edu_desc,
                "label": "学历门槛",
            },
        },
        "highlights": highlights,
        "suggestions": suggestions,
    }


# ── 批量入库（后台线程）────────────────────────────────────────────────

def _do_index(db_factory) -> None:

    global _index_status, _collection, _chroma_client

    with _lock:
        if _index_status["status"] == "indexing":
            return
        _index_status["status"] = "indexing"
       
        _collection = None
        _chroma_client = None

    try:
        from .models import Job
        import chromadb
        from chromadb.config import Settings

        db = db_factory()
        try:
            
            _CHROMA_PATH.mkdir(exist_ok=True)
            local_client = chromadb.PersistentClient(
                path=str(_CHROMA_PATH),
                settings=Settings(anonymized_telemetry=False),
            )
            local_col = local_client.get_or_create_collection(
                name=COLLECTION_NAME,
                metadata={"hnsw:space": "cosine"},
            )

            total = db.query(Job).count()
            _index_status["total"] = total

            
            existing = local_col.count()
            if existing >= int(total * 0.95):
                logger.info("向量库已有 %d 条，跳过入库", existing)
                _index_status.update({"status": "ready", "indexed": existing})
                return

           
            model = get_model()

            logger.info("开始批量入库：共 %d 个岗位（已入 %d）", total, existing)
            BATCH = 300
            offset = existing   
            while offset < total:
                batch = db.query(Job).order_by(Job.id).offset(offset).limit(BATCH).all()
                if not batch:
                    break
                texts = [build_job_text(j) for j in batch]
                ids = [str(j.id) for j in batch]
                embeddings = model.encode(
                    texts,
                    batch_size=64,
                    show_progress_bar=False,
                    normalize_embeddings=True,
                ).tolist()
                metadatas = [
                    {
                        "job_name": j.job_name or "",
                        "company_name": j.company_name or "",
                        "work_city": j.work_city or "",
                        "salary_min": j.salary_min or 0,
                        "salary_max": j.salary_max or 0,
                        "your_education": j.your_education or "",
                        "working_exp": j.working_exp or "",
                        "industry_name": j.industry_name or "",
                        "job_salary": j.job_salary or "",
                        "company_size": j.company_size or "",
                    }
                    for j in batch
                ]
                local_col.upsert(ids=ids, embeddings=embeddings, metadatas=metadatas)
                offset += len(batch)
                _index_status["indexed"] = offset
                logger.info("入库进度：%d / %d", offset, total)
        finally:
            db.close()

        _index_status.update({"status": "ready", "indexed": total})
        logger.info("向量库构建完成，共 %d 条", total)

    except Exception as e:
        logger.error("向量库入库失败：%s", e, exc_info=True)
        _index_status.update({"status": "error", "error": str(e)})


def start_indexing(db_factory) -> None:
    """在后台守护线程启动入库任务（幂等，已运行则跳过）"""
    with _lock:
        if _index_status["status"] in ("indexing", "ready"):
            return
    t = threading.Thread(target=_do_index, args=(db_factory,), daemon=True)
    t.start()


def get_index_status() -> dict:
    return dict(_index_status)



def auto_recommend(resume_text: str, top_k: int = 20, db=None) -> List[Dict]:
    from .models import Job as JobModel
    model = get_model()
    collection = get_collection()

    resume_vec = model.encode(
        resume_text,
        normalize_embeddings=True,
    ).tolist()

    candidates = min(top_k * 4, 100)
    results = collection.query(
        query_embeddings=[resume_vec],
        n_results=candidates,
        include=["metadatas", "distances"],
    )


    try:
        featured = collection.get(
            where={"featured": {"$eq": "1"}},
            include=["metadatas", "embeddings"],
        )
        if featured["ids"]:
            existing_ids = set(results["ids"][0])
            resume_vec_np = np.array(resume_vec)
            for fid, fmeta, femb in zip(
                featured["ids"], featured["metadatas"], featured["embeddings"]
            ):
                if fid not in existing_ids:
                    sim = float(np.dot(resume_vec_np, np.array(femb)))
                    dist_val = max(0.0, 1.0 - sim)
                    results["ids"][0].append(fid)
                    results["metadatas"][0].append(fmeta)
                    results["distances"][0].append(dist_val)
    except Exception as e:
        logger.debug("featured 岗位补充查询失败（可忽略）: %s", e)

    r_kw = _extract_keywords_set(resume_text, 30)


    all_candidate_ids = results["ids"][0]
    job_text_map: Dict[str, str] = {}   # job_id(str) -> full job text
    if db is not None:
        try:
            int_ids = [int(jid) for jid in all_candidate_ids]
            db_jobs = db.query(JobModel).filter(JobModel.id.in_(int_ids)).all()
            job_text_map = {str(j.id): build_job_text(j) for j in db_jobs}
        except Exception as e:
            logger.warning("auto_recommend: DB 批量查询失败，降级为 metadata 关键词: %s", e)

    matches = []
    for job_id, meta, dist in zip(
        results["ids"][0], results["metadatas"][0], results["distances"][0]
    ):
        # ── 语义分（40%）──────────────────────────────────────────────
        sbert_score = round((1.0 - dist) * 100, 1)
        if sbert_score >= 75:
            sem_desc = f"语义相似度 {sbert_score}%，简历方向与岗位高度吻合"
        elif sbert_score >= 55:
            sem_desc = f"语义相似度 {sbert_score}%，简历方向与岗位基本相关"
        else:
            sem_desc = f"语义相似度 {sbert_score}%，简历方向与岗位差异较大"


        if job_id in job_text_map:

            j_kw = _extract_keywords_set(job_text_map[job_id], 15)
        else:

            job_text_kw = (
                f"{meta.get('job_name','')} {meta.get('job_name','')} "
                f"{meta.get('industry_name','')} {meta.get('company_name','')}"
            )
            j_kw = _extract_keywords_set(job_text_kw, 12)
        matched_kws = sorted(r_kw & j_kw)
        raw_coverage = len(matched_kws) / max(len(j_kw), 1) if j_kw else 0.0

        skill_score = round(min(raw_coverage * 2.0 * 100, 100.0), 1)
        matched_count = len(matched_kws)
        total_jd_kw = max(len(j_kw), 1)

        exp_req = meta.get("working_exp", "") or ""
        exp_score_val, exp_desc = _exp_score(resume_text, exp_req)

        edu_req = meta.get("your_education", "") or ""
        edu_score_val, _, edu_meets = _edu_score(resume_text, edu_req)


        total_raw = 0.40 * sbert_score + 0.40 * skill_score + 0.20 * exp_score_val
        if not edu_meets:
            penalty = 0.80 if edu_score_val >= 60 else 0.65
            total_raw *= penalty
        total = round(max(0.0, min(100.0, total_raw)), 1)

        matches.append({
            "job_id": int(job_id),
            "match_score": total,
            **meta,
            "report": {
                "dimensions": {
                    "semantic": {
                        "score": sbert_score,
                        "label": "语义理解",
                        "match": sem_desc,
                    },
                    "skills": {
                        "score": skill_score,
                        "label": "技能匹配",
                        "match": f"命中岗位关键词 {matched_count}/{total_jd_kw} 个",
                        "matched_skills": matched_kws[:10],
                        "matched_count": matched_count,
                        "total_count": total_jd_kw,
                    },
                    "experience": {
                        "score": exp_score_val,
                        "label": "经验修正",
                        "match": exp_desc,
                    },
                },
                "edu_meets_threshold": edu_meets,
            },
        })
    return sorted(matches, key=lambda x: -x["match_score"])[:top_k]


def specific_match(resume_text: str, job) -> dict:

    return compute_match_report(resume_text, job)
